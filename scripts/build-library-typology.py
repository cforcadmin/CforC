# -*- coding: utf-8 -*-
"""
Χτίζει το έγγραφο «Τυπολογία Αρχειοθέτησης» της Ψηφιακής Βιβλιοθήκης.

ΓΙΑΤΙ ΠΑΡΑΓΕΤΑΙ ΚΑΙ ΔΕΝ ΓΡΑΦΕΤΑΙ: η πρώτη εκδοχή του εγγράφου είχε
πληκτρολογηθεί με το χέρι από την ιστοσελίδα. Στο μεταξύ η ταξινομία των
μελών προχώρησε, και το έγγραφο έμεινε πίσω — έλειπε ολόκληρη θεματική και
είχε δύο αόρατα λάθος γράμματα (λατινικό E, εισαγωγικό αντί για τόνο). Όσο
οι λίστες υπάρχουν σε δεύτερο αντίγραφο που συντηρείται χωριστά, η απόκλιση
είναι θέμα χρόνου. Εδώ δεν πληκτρολογεί κανείς τίποτα.

Τρέξιμο:
  npx tsx scripts/build-library-template.mjs /tmp/taxonomy.json
  python3 scripts/build-library-typology.py /tmp/taxonomy.json <φάκελος εξόδου>
"""
import json, sys, os, unicodedata
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def greek_upper(text):
    """Ελληνικά κεφαλαία: ο τόνος φεύγει, τα διαλυτικά μένουν.

    Το str.upper() της Python κρατά τον τόνο και βγάζει «ΕΚΠΑΊΔΕΥΣΗ» — λάθος
    στα ελληνικά. Ίδιος κανόνας με το toLocaleUpperCase('el') που χρησιμοποιεί
    η ιστοσελίδα.
    """
    d = unicodedata.normalize('NFD', text)
    d = ''.join(c for c in d if c not in ('\u0301', '\u0300', '\u0342', '\u0345'))
    return unicodedata.normalize('NFC', d).upper()

tx = json.load(open(sys.argv[1], encoding='utf-8'))
out = os.path.join(sys.argv[2], 'Τυπολογία Αρχειοθέτησης_CforC Digital Library (ενημερωμένη).docx')

# Το όνομα διαφέρει σκόπιμα από το πρωτότυπο της ομάδας. Και ό,τι κι αν
# γράφει εδώ, δεν πατάμε υπάρχον αρχείο χωρίς ρητή εντολή: το πρωτότυπο
# μπορεί να είναι η μόνη εκδοχή που έχει κάποιος.
if os.path.exists(out) and '--force' not in sys.argv:
    print(f'ΥΠΑΡΧΕΙ ΗΔΗ: {out}\nΔεν το πατάω. Ξανατρέξε με --force αν όντως το θέλεις.')
    sys.exit(1)

CORAL = RGBColor(0xE0, 0x6A, 0x45)
GREY = RGBColor(0x8A, 0x8A, 0x8A)
DARK = RGBColor(0x2D, 0x2D, 0x2D)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.4)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK

def para(text='', size=10.5, bold=False, color=DARK, space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p

# ── Τίτλος ──────────────────────────────────────────────────────
para('CforC Digital Library', 20, True, CORAL, space_after=0)
para('Τυπολογία Αρχειοθέτησης', 14, False, GREY, space_after=14)

para('Οι λίστες που ακολουθούν παράγονται αυτόματα από την ταξινομία της ιστοσελίδας '
     '— την ίδια που χρησιμοποιούν τα πεδία πρακτικής στα προφίλ των μελών. Έτσι η '
     'βιβλιοθήκη και το μητρώο μελών δεν μπορούν να αποκλίνουν.', 10, False, DARK, space_after=6)
para('Μην επεξεργάζεστε αυτό το αρχείο. Αν χρειάζεται προσθήκη ή διόρθωση, ζητήστε την '
     'και ενημερώνονται μαζί η ιστοσελίδα, το πρότυπο καταγραφής και το παρόν έγγραφο.',
     10, True, DARK, space_after=16)

# ── Θεματικές ───────────────────────────────────────────────────
para('Θεματικές και Υποθεματικές', 14, True, CORAL, space_before=6, space_after=2)
para(f"{len(tx['themes'])} θεματικές · {len(tx['pairs'])} υποθεματικές", 9, False, GREY, space_after=10)

by_theme = {}
for theme, sub in tx['pairs']:
    by_theme.setdefault(theme, []).append(sub)

for theme in tx['themes']:
    para(greek_upper(theme), 11, True, DARK, space_before=8, space_after=2)
    for sub in by_theme.get(theme, []):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(sub)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

# ── Είδη αρχείων ────────────────────────────────────────────────
doc.add_page_break()
para('Είδη αρχείων', 14, True, CORAL, space_after=2)
para(f"{len(tx['docTypes'])} είδη", 9, False, GREY, space_after=10)
for d in tx['docTypes']:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(d)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK

# ── Σημειώσεις καταγραφής ───────────────────────────────────────
para('Σημειώσεις καταγραφής', 14, True, CORAL, space_before=18, space_after=6)
notes = [
    ('Θεματική', 'Μία ανά αρχείο, από τη λίστα.'),
    ('Υποθεματική', 'Μία ή περισσότερες, χωρισμένες με κόμμα, από τις υποθεματικές της '
                    'επιλεγμένης θεματικής. Στην ιστοσελίδα η υποθεματική ακολουθεί '
                    'αυτόματα τη θεματική.'),
    ('Έτος κυκλοφορίας', 'Σκέτος αριθμός (π.χ. 2026), όχι κείμενο και όχι δεκαδικά.'),
    ('Σύνδεσμος πηγής', 'Η σελίδα του εκδότη. Είναι ο κύριος σύνδεσμος του αρχείου: '
                        'σωστή απόδοση, χωρίς ζήτημα πνευματικών δικαιωμάτων.'),
    ('Σύνδεσμος αρχείου', 'Το αντίγραφο στον φάκελο Assets_documents, ως αρχείο ασφαλείας '
                          'για όταν ο σύνδεσμος του εκδότη πάψει να δουλεύει.'),
    ('Δικαιώματα πρόσβασης', 'Ανεβάζουμε ΝΕΟ αντίγραφο στον φάκελο· δεν μετακινούμε αρχείο '
                             'που υπάρχει ήδη αλλού στο Drive, γιατί κουβαλά μαζί του τον '
                             'παλιό του διαμοιρασμό. Κάθε αρχείο μένει σε «Περιορισμένη '
                             'πρόσβαση» — την πρόσβαση των μελών την αναλαμβάνει η '
                             'ιστοσελίδα.'),
]
for label, text in notes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + ' — ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK

f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.paragraph_format.space_before = Pt(24)
r = f.add_run('Παράγεται από την ταξινομία της ιστοσελίδας · Culture for Change')
r.font.size = Pt(8)
r.font.color.rgb = GREY

doc.save(out)
print('γράφτηκε:', out)
